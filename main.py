import logging
import re
import os
from io import BytesIO
from datetime import datetime

from telegram import Update, InputFile
from telegram.ext import (
    Application, CommandHandler, MessageHandler, ConversationHandler,
    ContextTypes, filters
)
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ✅ Add your bot token here
TOKEN = os.getenv("BOT_TOKEN")

if not TOKEN:
    raise ValueError("No BOT_TOKEN found in environment variables")

application = Application.builder().token(TOKEN).build()

logging.basicConfig(
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s", level=logging.INFO
)
logger = logging.getLogger(__name__)

# --- Conversation states ---
(
    CLASS_INSTRUCTOR, SPEAKING_EXAMINER, INVIGILATOR, LEVEL,
    TEST_DATE, FINISH_DATE, STUDENT_COUNT, SKILLS_COUNT,
    SKILL_NAME, SKILL_MAX_SCORE, STUDENT_DETAILS, SCORES
) = range(12)

DATE_REGEX = r"^\d{4}-\d{2}-\d{2}$"

# --- Handlers ---
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 Hello! Let's start by collecting test details.\n\nWho is the class instructor?"
    )
    return CLASS_INSTRUCTOR

async def get_class_instructor(update: Update, context):
    context.user_data['class_instructor'] = update.message.text.strip()
    await update.message.reply_text("Who is the speaking examiner?")
    return SPEAKING_EXAMINER

async def get_speaking_examiner(update: Update, context):
    context.user_data['speaking_examiner'] = update.message.text.strip()
    await update.message.reply_text("Who is the invigilator?")
    return INVIGILATOR

async def get_invigilator(update: Update, context):
    context.user_data['invigilator'] = update.message.text.strip()
    await update.message.reply_text("What is the level?")
    return LEVEL

async def get_level(update: Update, context):
    context.user_data['level'] = update.message.text.strip()
    await update.message.reply_text("What is the test date? (Format: YYYY-MM-DD)")
    return TEST_DATE

async def get_test_date(update: Update, context):
    text = update.message.text.strip()
    if not re.match(DATE_REGEX, text):
        await update.message.reply_text("Invalid date format. Please enter the test date as YYYY-MM-DD.")
        return TEST_DATE
    try:
        datetime.strptime(text, "%Y-%m-%d")
    except ValueError:
        await update.message.reply_text("Invalid date. Please enter a real date in YYYY-MM-DD format.")
        return TEST_DATE
    context.user_data['test_date'] = text
    await update.message.reply_text("What is the finish date? (Format: YYYY-MM-DD)")
    return FINISH_DATE

async def get_finish_date(update: Update, context):
    text = update.message.text.strip()
    if not re.match(DATE_REGEX, text):
        await update.message.reply_text("Invalid date format. Please enter the finish date as YYYY-MM-DD.")
        return FINISH_DATE
    try:
        datetime.strptime(text, "%Y-%m-%d")
    except ValueError:
        await update.message.reply_text("Invalid date. Please enter a real date in YYYY-MM-DD format.")
        return FINISH_DATE
    context.user_data['finish_date'] = text
    await update.message.reply_text("How many students do you want to enter?")
    return STUDENT_COUNT

async def get_student_count(update: Update, context):
    text = update.message.text.strip()
    if not text.isdigit() or int(text) <= 0:
        await update.message.reply_text("Please enter a valid positive number for students count.")
        return STUDENT_COUNT
    context.user_data['student_count'] = int(text)
    context.user_data['students'] = []
    context.user_data['current_student'] = 1
    await update.message.reply_text("How many skills does the group have?")
    return SKILLS_COUNT

async def get_skills_count(update: Update, context):
    text = update.message.text.strip()
    if not text.isdigit() or int(text) <= 0:
        await update.message.reply_text("Please enter a valid positive number of skills.")
        return SKILLS_COUNT
    context.user_data['skills_count'] = int(text)
    context.user_data['skills'] = []
    context.user_data['current_skill'] = 1
    await update.message.reply_text("Enter the name of skill 1:")
    return SKILL_NAME

async def get_skill_name(update: Update, context):
    skill_name = update.message.text.strip()
    context.user_data['current_skill_name'] = skill_name
    await update.message.reply_text(f"Enter the max score for '{skill_name}':")
    return SKILL_MAX_SCORE

async def get_skill_max_score(update: Update, context):
    text = update.message.text.strip()
    try:
        max_score = float(text)
        if max_score <= 0:
            raise ValueError
    except ValueError:
        await update.message.reply_text("Please enter a valid positive number for max score.")
        return SKILL_MAX_SCORE
    skill = {
        'name': context.user_data['current_skill_name'],
        'max_score': max_score
    }
    context.user_data['skills'].append(skill)

    if context.user_data['current_skill'] < context.user_data['skills_count']:
        context.user_data['current_skill'] += 1
        await update.message.reply_text(f"Enter the name of skill {context.user_data['current_skill']}:")
        return SKILL_NAME
    else:
        await update.message.reply_text("Enter the full name of student 1:")
        return STUDENT_DETAILS

async def get_student_details(update: Update, context):
    student_name = update.message.text.strip()
    context.user_data['current_student_name'] = student_name
    context.user_data['scores'] = []
    context.user_data['current_skill_score'] = 1
    await update.message.reply_text(
        f"Enter the score for skill 1 '{context.user_data['skills'][0]['name']}' for student {student_name}:"
    )
    return SCORES

async def get_scores(update: Update, context):
    text = update.message.text.strip()
    skill_idx = context.user_data['current_skill_score'] - 1
    max_score = context.user_data['skills'][skill_idx]['max_score']
    try:
        score = float(text)
        if score < 0 or score > max_score:
            raise ValueError
    except ValueError:
        await update.message.reply_text(f"Please enter a valid score between 0 and {max_score}.")
        return SCORES
    context.user_data['scores'].append(score)

    if context.user_data['current_skill_score'] < context.user_data['skills_count']:
        context.user_data['current_skill_score'] += 1
        skill_name = context.user_data['skills'][context.user_data['current_skill_score'] - 1]['name']
        await update.message.reply_text(
            f"Enter the score for skill {context.user_data['current_skill_score']} '{skill_name}' "
            f"for student {context.user_data['current_student_name']}:"
        )
        return SCORES
    else:
        student_data = {
            'name': context.user_data['current_student_name'],
            'scores': context.user_data['scores'].copy(),
            'max_scores': [s['max_score'] for s in context.user_data['skills']]
        }
        context.user_data['students'].append(student_data)
        context.user_data['scores'].clear()

        if context.user_data['current_student'] < context.user_data['student_count']:
            context.user_data['current_student'] += 1
            await update.message.reply_text(f"Enter the full name of student {context.user_data['current_student']}:")
            return STUDENT_DETAILS
        else:
            await generate_word_report(update, context)
            return ConversationHandler.END

async def generate_word_report(update: Update, context):
    doc = Document()

    # Optional logo
    try:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        image_path = os.path.join(current_dir, "images.jpg")

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.5))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    except Exception as e:
        logger.warning(f"Could not add logo image: {e}")

    # Title
    title = doc.add_heading('End-of-Course Test Results Form', level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Info table
    info_table = doc.add_table(rows=2, cols=3)
    info_table.style = 'Table Grid'
    labels = ['Class Instructor', 'Finish Date', 'Level',
              'Test Date', 'Invigilator', 'Speaking Examiner']
    values = [
        context.user_data['class_instructor'],
        context.user_data['finish_date'],
        context.user_data['level'],
        context.user_data['test_date'],
        context.user_data['invigilator'],
        context.user_data['speaking_examiner']
    ]

    for i in range(2):
        for j in range(3):
            cell = info_table.cell(i, j)
            label = labels[i * 3 + j]
            value = values[i * 3 + j]
            cell.text = f"{label}\n{value}"
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            if para.runs:
                run = para.runs[0]
            else:
                run = para.add_run()
            run.font.size = Pt(10)

    doc.add_paragraph()

    # Skills and students
    skills = context.user_data['skills']
    skill_names = [s['name'] for s in skills]
    skill_maxes = [s['max_score'] for s in skills]
    num_skills = len(skills)
    students = context.user_data['students']

    # Calculate percentages and pass/fail
    for st in students:
        total_score = sum(st['scores'])
        total_max = sum(s['max_score'] for s in skills)
        st['percent'] = round(total_score / total_max * 100, 2) if total_max > 0 else 0
        st['result'] = "Pass" if st['percent'] >= 60 else "Fail"

    students.sort(key=lambda s: s['percent'], reverse=True)

    # Create table
    table = doc.add_table(rows=2, cols=2 + num_skills + 2)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Names"
    for i, name in enumerate(skill_names):
        hdr_cells[2 + i].text = name
    hdr_cells[2 + num_skills].text = "Overall (%)"
    hdr_cells[2 + num_skills + 1].text = "Pass/Fail"

    # Sub-header row
    sub_hdr_cells = table.rows[1].cells
    sub_hdr_cells[0].text = ""
    sub_hdr_cells[1].text = ""
    for i, max_score in enumerate(skill_maxes):
        sub_hdr_cells[2 + i].text = str(max_score)
    sub_hdr_cells[2 + num_skills].text = "100%"
    sub_hdr_cells[2 + num_skills + 1].text = ""

    # Add student rows
    for idx, st in enumerate(students, start=1):
        row_cells = table.add_row().cells
        data = [str(idx), st['name']] + [str(s) for s in st['scores']] + [str(st['percent']), st['result']]
        for k, cell in enumerate(row_cells):
            cell.text = data[k]
            para = cell.paragraphs[0]
            para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if k != 1 else WD_PARAGRAPH_ALIGNMENT.LEFT
            if para.runs:
                run = para.runs[0]
            else:
                run = para.add_run()
            run.font.size = Pt(10)
            if k == len(row_cells) - 1:  # Last column Pass/Fail color
                run.font.color.rgb = RGBColor(0, 128, 0) if st['result'] == "Pass" else RGBColor(255, 0, 0)

    # Save document to memory
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    # Send the document
    await update.message.reply_document(
        document=InputFile(bio, filename="End-of-Course Test Results.docx")
    )
    await update.message.reply_text("✅ Test results document created successfully.")

async def cancel(update: Update, context):
    await update.message.reply_text('Operation cancelled. You can restart with /start.')
    return ConversationHandler.END

# --- Conversation handler setup ---
conv_handler = ConversationHandler(
    entry_points=[CommandHandler("start", start)],
    states={
        CLASS_INSTRUCTOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_class_instructor)],
        SPEAKING_EXAMINER: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_speaking_examiner)],
        INVIGILATOR: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_invigilator)],
        LEVEL: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_level)],
        TEST_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_test_date)],
        FINISH_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_finish_date)],
        STUDENT_COUNT: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_student_count)],
        SKILLS_COUNT: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_skills_count)],
        SKILL_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_skill_name)],
        SKILL_MAX_SCORE: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_skill_max_score)],
        STUDENT_DETAILS: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_student_details)],
        SCORES: [MessageHandler(filters.TEXT & ~filters.COMMAND, get_scores)],
    },
    fallbacks=[CommandHandler("cancel", cancel)],
    allow_reentry=True,
)
application.add_handler(conv_handler)

if __name__ == "__main__":
    PORT = int(os.environ.get("PORT", 10000))
    RENDER_EXTERNAL_URL = os.environ.get("RENDER_EXTERNAL_URL")

    if not RENDER_EXTERNAL_URL:
        raise ValueError("RENDER_EXTERNAL_URL not set")

    print("🤖 Bot is starting with webhook...")

    application.run_webhook(
        listen="0.0.0.0",
        port=PORT,
        webhook_url=f"{RENDER_EXTERNAL_URL}/webhook",
    )
